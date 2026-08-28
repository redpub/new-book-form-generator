import sys
from docx import Document
from lxml import etree

def inspect_doc(path):
    print(f"=== Inspecting {path} ===")
    doc = Document(path)
    
    # 1. Paragraphs (non-table)
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    
    # Check headers/footers
    sections = doc.sections
    print(f"Total Sections: {len(sections)}")
    for i, s in enumerate(sections):
        print(f"Section {i}:")
        for h_type in ['header', 'first_page_header', 'even_page_header']:
            h = getattr(s, h_type, None)
            if h and h.paragraphs:
                txt = "".join(p.text for p in h.paragraphs).strip()
                if txt:
                    print(f"  {h_type}: {txt[:100]}")
        for f_type in ['footer', 'first_page_footer', 'even_page_footer']:
            f = getattr(s, f_type, None)
            if f and f.paragraphs:
                txt = "".join(p.text for p in f.paragraphs).strip()
                if txt:
                    print(f"  {f_type}: {txt[:100]}")

    # 2. Tables
    print(f"Total Tables: {len(doc.tables)}")
    for t_idx, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.columns) if table.rows else 0
        print(f"Table {t_idx}: {rows} rows x {cols} columns")
        # Let's print some row/cell structure or specific text
        for r_idx, row in enumerate(table.rows[:5]): # first few rows
            cell_texts = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            # remove duplicates for merged cells if we want, but let's just inspect them
            print(f"  Row {r_idx}: {cell_texts[:6]}")
        if len(table.rows) > 5:
            print(f"  ... and {len(table.rows)-5} more rows")

inspect_doc("template.docx")
inspect_doc("template-v2.docx")
