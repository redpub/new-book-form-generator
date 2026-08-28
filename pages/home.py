from html.parser import HTMLParser as _HTMLParser
import io
import json
import re
import time
from typing import Any

from json_repair import repair_json
import streamlit as st
from docx import Document
from docx.oxml.ns import qn
from docxtpl import DocxTemplate
from pydantic import BaseModel, Field, ValidationError
from tenacity import retry, stop_after_attempt, wait_exponential

from app_config import APP_DIR, PROMPT_FILE, TEMPLATE_FILE, VERTEX_MODEL
from auth import enforce_workspace_auth
from debug_store import add_run
from field_config import CHECKBOX_CHAR, UNCHECKED_CHAR, SECTIONS
from streamlit_quill import st_quill


# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------

class BookFormData(BaseModel):
    title: str | None = ""
    title_chi: str | None = ""
    title_eng: str | None = ""
    subtitle: str | None = ""
    series: str | None = ""
    author_name: str | None = ""
    author_bio: str | None = ""
    target_audience: str | None = ""
    synopsis: str | None = ""
    isbn: str | None = ""
    publication_date: str | None = ""
    language: list[str] | str | None = Field(default_factory=list)
    category: str | None = ""
    keywords: list[str] = Field(default_factory=list)
    page_count: str | None = ""
    trim_size: str | None = ""
    thickness: str | None = ""
    binding: str | None = ""
    print_color: str | None = ""
    price: str | None = ""
    price_hkd: str | None = ""
    price_twd: str | None = ""
    contributor: str | None = ""
    editor_notes: str | None = ""
    publisher: str | None = ""
    book_highlights: str | None = ""
    endorsements: str | None = ""
    table_of_contents: str | None = ""
    extras: dict[str, Any] = Field(default_factory=dict)


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def read_prompt_text() -> str:
    if not PROMPT_FILE.exists():
        raise FileNotFoundError(
            "找不到專案根目錄的 prompt.txt，請先建立後再執行擷取。"
        )
    return PROMPT_FILE.read_text(encoding="utf-8").strip()


def _para_to_formatted(para) -> str:
    """Return paragraph text with **bold**, *italic*, __underline__ run markers."""
    parts: list[str] = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        if run.underline:
            text = f"__{text}__"
        if run.italic:
            text = f"*{text}*"
        if run.bold:
            text = f"**{text}**"
        parts.append(text)
    return "".join(parts)


def _get_list_type(para) -> str:
    """Return 'bullet', 'number', or '' (not a list paragraph)."""
    style_name = (para.style.name or "") if para.style else ""
    style_lower = style_name.lower()
    pPr = para._p.find(qn("w:pPr"))
    num_pr = pPr.find(qn("w:numPr")) if pPr is not None else None
    if num_pr is None and "list" not in style_lower:
        return ""
    if "number" in style_lower:
        return "number"
    return "bullet"


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    lines: list[str] = []
    num_counters: dict[str, int] = {}

    for para in doc.paragraphs:
        content = _para_to_formatted(para)
        if not content.strip():
            continue
        list_type = _get_list_type(para)
        if list_type == "bullet":
            lines.append(f"• {content}")
        elif list_type == "number":
            pPr = para._p.find(qn("w:pPr"))
            num_pr = pPr.find(qn("w:numPr")) if pPr is not None else None
            num_id = "0"
            if num_pr is not None:
                num_id_el = num_pr.find(qn("w:numId"))
                if num_id_el is not None:
                    num_id = num_id_el.get(qn("w:val"), "0")
            num_counters[num_id] = num_counters.get(num_id, 0) + 1
            lines.append(f"{num_counters[num_id]}. {content}")
        else:
            lines.append(content)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    lines.append(text)

    for section in doc.sections:
        for para in section.header.paragraphs + section.footer.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)

    return "\n".join(lines)


def flatten_extracted_data(payload: dict[str, Any]) -> dict[str, str]:
    flat: dict[str, str] = {}

    def walk(obj: Any, prefix: str = "") -> None:
        if isinstance(obj, dict):
            for key, value in obj.items():
                walk(value, f"{prefix}.{key}" if prefix else key)
        elif isinstance(obj, list):
            flat[prefix] = (
                ", ".join(obj)
                if all(isinstance(i, str) for i in obj)
                else json.dumps(obj, ensure_ascii=False)
            )
        elif obj is None:
            flat[prefix] = ""
        else:
            flat[prefix] = str(obj)

    walk(payload)
    return flat


# ---------------------------------------------------------------------------
# Google Vertex helpers
# ---------------------------------------------------------------------------

def parse_json_from_text(raw_text: str) -> dict[str, Any]:
    fenced = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", raw_text, re.DOTALL)
    candidate = fenced.group(1) if fenced else raw_text

    # 1. Try strict parse first
    try:
        parsed = json.loads(candidate)
        if isinstance(parsed, dict):
            return parsed
    except json.JSONDecodeError:
        pass

    # 2. LLM output often has unescaped quotes inside string values;
    #    json_repair fixes these automatically.
    try:
        repaired = repair_json(candidate, return_objects=True)
        if isinstance(repaired, dict) and repaired:
            return repaired
    except Exception:
        pass

    # 3. Last resort: brace-extract then repair
    brace = re.search(r"\{[\s\S]*\}", candidate)
    if brace:
        try:
            repaired = repair_json(brace.group(0), return_objects=True)
            if isinstance(repaired, dict) and repaired:
                return repaired
        except Exception:
            pass

    raise ValueError("無法從 Google Vertex 回應中解析 JSON 物件。")


def normalize_extracted_payload(payload: dict[str, Any]) -> dict[str, Any]:
    known_keys = {
        "title", "title_chi", "title_eng", "subtitle", "series", "author_name", "author_bio",
        "target_audience", "synopsis", "isbn", "publication_date",
        "language", "category", "keywords", "page_count", "trim_size", "thickness",
        "binding", "print_color", "price", "price_hkd", "price_twd", "contributor", "editor_notes",
        "publisher", "book_highlights", "endorsements", "table_of_contents", "extras",
    }
    extras: dict[str, Any] = payload.get("extras") or {}
    for key, value in payload.items():
        if key not in known_keys:
            extras[key] = value
    data = dict(payload)
    for price_key in ("price_hkd", "price_twd"):
        if price_key in data:
            data[price_key] = _sanitize_non_negative_integer(data[price_key])
    data["extras"] = extras if isinstance(extras, dict) else {}
    return BookFormData.model_validate(data).model_dump()


def _collect_text_candidates(obj: Any) -> list[str]:
    if isinstance(obj, str):
        return [obj]
    if isinstance(obj, dict):
        return [c for v in obj.values() for c in _collect_text_candidates(v)]
    if isinstance(obj, list):
        return [c for item in obj for c in _collect_text_candidates(item)]
    return []


def build_field_reference() -> str:
    """Build the JSON schema + label block substituted into the {FIELD_REFERENCE} placeholder."""
    top_level: dict[str, Any] = {}
    label_lines: list[str] = []
    add_trim_size = False
    add_language = False
    language_labels: list[str] = []

    for section in SECTIONS:
        for field in section["fields"]:
            if not field.get("included_in_ai_prompt"):
                continue
            src = field.get("source_key", "")
            label = field.get("label", "")
            if not src:
                continue
            if src.startswith("#lang:"):
                add_language = True
                if label not in language_labels:
                    language_labels.append(label)
                continue
            if src.startswith("#"):
                add_trim_size = True
                label_lines.append(
                    f'  {label} \u2192 trim_size (parse from e.g. "150mm X 210mm")')
                continue
            top_level[src] = ""
            label_lines.append(f"  {label} \u2192 {src}")

    if add_language:
        top_level["language"] = []
        opts_str = "\u3001".join(language_labels)
        label_lines.append(
            f"  \u8a9e\u7a2e \u2192 language (array, use applicable values from: {opts_str})")
    if add_trim_size:
        top_level["trim_size"] = ""

    schema_json = json.dumps(top_level, indent=2, ensure_ascii=False)
    label_block = "\n".join(
        ["Form field labels (Chinese label \u2192 JSON key):"] + label_lines)
    return f"{schema_json}\n\n{label_block}"


def compose_vertex_prompt(prompt_text: str, document_text: str) -> str:
    merged = prompt_text.replace("{FIELD_REFERENCE}", build_field_reference())
    return (
        f"{merged}\n\n"
        "INPUT DOCUMENT:\n"
        f"{document_text}\n\n"
        "Return only one valid JSON object."
    )


def _parse_google_service_account_info(raw_service_account: Any) -> dict[str, Any]:
    """Parse and validate service-account JSON from Streamlit secrets."""
    if isinstance(raw_service_account, str):
        try:
            service_account_info = json.loads(raw_service_account)
        except json.JSONDecodeError as exc:
            raise ValueError(
                "google_vertex.service_account_json 不是有效的 JSON。"
            ) from exc
    elif isinstance(raw_service_account, dict):
        service_account_info = raw_service_account
    elif hasattr(raw_service_account, "items"):
        service_account_info = dict(raw_service_account.items())
    else:
        raise ValueError(
            "google_vertex.service_account_json 格式錯誤，需為 JSON 字串或物件。"
        )

    required = ("client_email", "token_uri", "private_key")
    missing = [key for key in required if not service_account_info.get(key)]
    if missing:
        raise ValueError(
            "Google service account 缺少必要欄位：" + ", ".join(missing)
        )
    return service_account_info


def get_google_vertex_client() -> Any:
    """Create a Vertex AI client from the [google_vertex] secrets section."""
    try:
        from google import genai
        from google.oauth2 import service_account
    except ImportError as exc:
        raise ValueError(
            "缺少 Google Vertex 相依套件，請安裝 google-genai 與 google-auth。"
        ) from exc

    settings = st.secrets.get("google_vertex", {})
    project_id = settings.get("project_id")
    location = settings.get("location")
    raw_service_account = settings.get("service_account_json")
    if not project_id or not location or not raw_service_account:
        raise ValueError(
            "請在 secrets.toml 的 [google_vertex] 提供 project_id、location、"
            "service_account_json。"
        )

    credentials = service_account.Credentials.from_service_account_info(
        _parse_google_service_account_info(raw_service_account),
        scopes=["https://www.googleapis.com/auth/cloud-platform"],
    )
    return genai.Client(
        vertexai=True,
        project=project_id,
        location=location,
        credentials=credentials,
    )


@retry(wait=wait_exponential(multiplier=1, min=1, max=8), stop=stop_after_attempt(3), reraise=True)
def call_google_vertex(client: Any, model: str, composed_prompt: str) -> tuple[dict[str, Any], str]:
    """Call Gemini on Vertex AI and return (parsed JSON, raw response text)."""
    response = client.models.generate_content(
        model=model,
        contents=composed_prompt,
        config={
            "temperature": 0,
            "response_mime_type": "application/json",
        },
    )
    raw_text = response.text or ""
    if not raw_text:
        raise RuntimeError("Google Vertex 回傳空白回應。")
    return parse_json_from_text(raw_text), raw_text


# ---------------------------------------------------------------------------
# Form value helpers (config-driven)
# ---------------------------------------------------------------------------

def _parse_trim_dimension(trim_size: str, index: int) -> str:
    """Parse width (index=0) or height (index=1) from e.g. '150mm X 210mm'."""
    parts = re.split(r"\s*[xX×]\s*", trim_size.strip())
    if len(parts) > index:
        return re.sub(r"[^\d.]", "", parts[index])
    return ""


_DATE_FORMATS = [
    "%Y-%m-%d", "%d/%m/%Y", "%d/%m/%y", "%Y/%m/%d",
    "%d-%m-%Y", "%d-%m-%y", "%B %d, %Y", "%b %d, %Y",
    "%d %B %Y", "%d %b %Y", "%Y年%m月%d日",
]


def _format_date_ddmmyy(value: str) -> str:
    """Try to parse *value* and return dd/mm/yy; return original on failure."""
    from datetime import datetime
    v = value.strip()
    for fmt in _DATE_FORMATS:
        try:
            return datetime.strptime(v, fmt).strftime("%d/%m/%y")
        except ValueError:
            continue
    return v


def _parse_to_date(value: str):
    """Parse *value* with known formats and return a datetime.date, or None."""
    from datetime import datetime
    v = value.strip()
    for fmt in _DATE_FORMATS:
        try:
            return datetime.strptime(v, fmt).date()
        except ValueError:
            continue
    return None


_CM_PLACEHOLDERS = {"width", "height", "thickness"}


def _mm_to_cm(value: str) -> str:
    """Convert a numeric mm string to cm (divide by 10), e.g. '150' → '15'."""
    try:
        cm = float(value) / 10
        return str(int(cm)) if cm == int(cm) else f"{cm:.1f}"
    except (ValueError, TypeError):
        return value


# ---------------------------------------------------------------------------
# HTML helpers (for st_quill → DOCX plain-text conversion)
# ---------------------------------------------------------------------------


class _QuillHTMLParser(_HTMLParser):
    """Convert st_quill HTML output to plain text preserving list structure."""

    def __init__(self) -> None:
        super().__init__()
        self._parts: list[str] = []
        self._list_stack: list[str] = []
        self._li_counters: dict[int, int] = {}

    # type: ignore[override]
    def handle_starttag(self, tag: str, attrs) -> None:
        if tag == "ul":
            self._list_stack.append("ul")
        elif tag == "ol":
            self._list_stack.append("ol")
            self._li_counters[len(self._list_stack) - 1] = 0
        elif tag == "li":
            if self._list_stack and self._list_stack[-1] == "ol":
                depth = len(self._list_stack) - 1
                self._li_counters[depth] = self._li_counters.get(depth, 0) + 1
                self._parts.append(f"\n{self._li_counters[depth]}. ")
            else:
                self._parts.append("\n• ")
        elif tag in ("p", "div", "br"):
            self._parts.append("\n")

    def handle_endtag(self, tag: str) -> None:  # type: ignore[override]
        if tag in ("ul", "ol") and self._list_stack:
            self._list_stack.pop()

    def handle_data(self, data: str) -> None:  # type: ignore[override]
        self._parts.append(data)

    def get_text(self) -> str:
        raw = "".join(self._parts)
        return re.sub(r"\n{3,}", "\n\n", raw).strip()


def _html_to_plain_text(html_str: str) -> str:
    """Convert st_quill HTML to plain text suitable for DOCX template substitution."""
    if not html_str:
        return ""
    parser = _QuillHTMLParser()
    parser.feed(html_str)
    return parser.get_text()


# ---------------------------------------------------------------------------
# Bullet-list helpers for native Word bullet formatting
# ---------------------------------------------------------------------------

# Sentinels inserted into the docxtpl context; post-processing replaces them
# with properly-formatted Word bullet paragraphs.
_BULLET_SENTINEL: dict[str, str] = {
    "bookIntro": "XXBULLETINTROXX",
    "bookHighlights": "XXBULLETHIGHLIGHTSXX",
}


class _HTMLItemParser(_HTMLParser):
    """Parse st_quill HTML into a list of (style, text) tuples.

    style values:
      'bullet'   – <ul><li> item → Word native bullet paragraph
      'number:N' – <ol><li> item (N is the 1-based counter)
      'normal'   – <p> content   → Normal paragraph
    """

    def __init__(self) -> None:
        super().__init__()
        self._items: list[tuple[str, str]] = []
        self._buf: list[str] = []
        self._list_stack: list[str] = []
        self._ol_counters: dict[int, int] = {}
        self._in_li = False
        self._style = "normal"

    # type: ignore[override]
    def handle_starttag(self, tag: str, attrs) -> None:
        if tag == "ul":
            self._list_stack.append("ul")
        elif tag == "ol":
            self._list_stack.append("ol")
            self._ol_counters[len(self._list_stack) - 1] = 0
        elif tag == "li":
            self._in_li = True
            self._buf = []
            depth = len(self._list_stack) - 1
            if self._list_stack and self._list_stack[-1] == "ol":
                self._ol_counters[depth] = self._ol_counters.get(depth, 0) + 1
                self._style = f"number:{self._ol_counters[depth]}"
            else:
                self._style = "bullet"
        elif tag == "p" and not self._in_li:
            self._buf = []
            self._style = "normal"
        elif tag == "br":
            self._buf.append("\n")

    def handle_endtag(self, tag: str) -> None:  # type: ignore[override]
        if tag in ("ul", "ol") and self._list_stack:
            self._list_stack.pop()
        elif tag == "li":
            text = "".join(self._buf).strip()
            if text:
                self._items.append((self._style, text))
            self._in_li = False
            self._buf = []
        elif tag == "p" and not self._in_li:
            text = "".join(self._buf).strip()
            if text:
                self._items.append(("normal", text))
            self._buf = []

    def handle_data(self, data: str) -> None:  # type: ignore[override]
        self._buf.append(data)

    def get_items(self) -> list[tuple[str, str]]:
        return self._items


def _parse_html_to_items(html_str: str) -> list[tuple[str, str]]:
    """Parse quill HTML into (style, text) tuples; falls back to plain text if needed."""
    if not html_str:
        return []
    parser = _HTMLItemParser()
    parser.feed(html_str)
    items = parser.get_items()
    if not items:
        plain = _html_to_plain_text(html_str)
        if plain:
            return [("normal", plain)]
    return items


def _find_bullet_numid(doc) -> str:
    """Return a numId string from the document's numbering part for bullet lists.

    Prefers Symbol \\uf0b7 or Unicode • bullets; falls back to any bullet numId.
    Returns empty string if no numbering part exists.
    """
    try:
        from docx.oxml.ns import qn as _qn
        num_xml = doc.part.numbering_part._element
    except AttributeError:
        return ""

    preferred = {"\uf0b7", "\u2022", "•"}
    fallback = ""

    # Build a quick lookup: abstractNumId → (fmt, lvlText)
    abs_info: dict[str, tuple[str, str]] = {}
    for abs_num in num_xml.findall(_qn("w:abstractNum")):
        aid = abs_num.get(_qn("w:abstractNumId"), "")
        lvl = abs_num.find(f'.//{_qn("w:lvl")}[@{_qn("w:ilvl")}="0"]')
        if lvl is None:
            continue
        nf = lvl.find(_qn("w:numFmt"))
        lt = lvl.find(_qn("w:lvlText"))
        fmt = nf.get(_qn("w:val"), "") if nf is not None else ""
        txt = lt.get(_qn("w:val"), "") if lt is not None else ""
        abs_info[aid] = (fmt, txt)

    for num in num_xml.findall(_qn("w:num")):
        nid = num.get(_qn("w:numId"), "")
        ref = num.find(_qn("w:abstractNumId"))
        if ref is None:
            continue
        aid = ref.get(_qn("w:val"), "")
        fmt, txt = abs_info.get(aid, ("", ""))
        if fmt != "bullet":
            continue
        if not fallback:
            fallback = nid
        if txt in preferred:
            return nid

    return fallback


def _apply_bullet_formatting(
    doc_bytes: bytes,
    sentinel_to_items: dict[str, list[tuple[str, str]]],
) -> bytes:
    """Replace sentinel placeholders in a rendered DOCX with Word bullet paragraphs.

    For each sentinel found in a table cell, the sentinel paragraph is removed and
    replaced with one paragraph per item:
    - style='bullet'   → paragraph with <w:numPr> pointing to a bullet numId
    - style='number:N' → plain Normal paragraph prefixed with "N. "
    - style='normal'   → plain Normal paragraph
    """
    if not sentinel_to_items:
        return doc_bytes

    import copy
    from docx import Document as _DocX
    from docx.oxml import OxmlElement as _OxmlEl
    from docx.oxml.ns import qn as _qn

    doc = _DocX(io.BytesIO(doc_bytes))
    bullet_numid = _find_bullet_numid(doc)

    _XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"

    for table in doc.tables:
        # Iterate physical tc elements from the table XML to avoid
        # python-docx's merged-cell aliasing that makes the same tc
        # appear under multiple (row, col) positions.
        replaced: set[str] = set()
        for tc in table._tbl.findall(f'.//{_qn("w:tc")}'):
            if len(replaced) == len(sentinel_to_items):
                break  # all sentinels replaced
            # Gather w:p elements directly from tc to avoid python-docx wrappers
            for para_el in list(tc.findall(_qn("w:p"))):
                runs = para_el.findall(f'.//{_qn("w:r")}')
                para_text = "".join(
                    (t.text or "")
                    for r in runs
                    for t in r.findall(_qn("w:t"))
                )
                matched = next(
                    (s for s in sentinel_to_items
                     if s not in replaced and s in para_text),
                    None,
                )
                if matched is None:
                    continue

                items = sentinel_to_items[matched]
                insert_idx = list(tc).index(para_el)

                # Grab run rPr for font inheritance (may be None)
                orig_rpr = para_el.find(f".//{_qn('w:rPr')}")

                tc.remove(para_el)

                for offset, (style, text) in enumerate(items):
                    new_p = _OxmlEl("w:p")

                    # Paragraph properties
                    pPr = _OxmlEl("w:pPr")
                    if style == "bullet" and bullet_numid:
                        numPr = _OxmlEl("w:numPr")
                        ilvl = _OxmlEl("w:ilvl")
                        ilvl.set(_qn("w:val"), "0")
                        numId_el = _OxmlEl("w:numId")
                        numId_el.set(_qn("w:val"), bullet_numid)
                        numPr.append(ilvl)
                        numPr.append(numId_el)
                        pPr.append(numPr)
                    new_p.append(pPr)

                    # Run
                    run_text = text
                    if style.startswith("number:"):
                        n = style.split(":", 1)[1]
                        run_text = f"{n}. {text}"

                    r = _OxmlEl("w:r")
                    if orig_rpr is not None:
                        r.append(copy.deepcopy(orig_rpr))
                    t = _OxmlEl("w:t")
                    t.text = run_text
                    t.set(_XML_SPACE, "preserve")
                    r.append(t)
                    new_p.append(r)

                    tc.insert(insert_idx + offset, new_p)

                replaced.add(matched)
                break  # one sentinel per tc
                runs = para_el.findall(f'.//{_qn("w:r")}')
                para_text = "".join(
                    (t.text or "")
                    for r in runs
                    for t in r.findall(_qn("w:t"))
                )
                matched = next(
                    (s for s in sentinel_to_items
                     if s not in replaced and s in para_text),
                    None,
                )
                if matched is None:
                    continue

                items = sentinel_to_items[matched]
                insert_idx = list(tc).index(para_el)

                # Grab run rPr for font inheritance (may be None)
                orig_rpr = para_el.find(f".//{_qn('w:rPr')}")

                tc.remove(para_el)

                for offset, (style, text) in enumerate(items):
                    new_p = _OxmlEl("w:p")

                    # Paragraph properties
                    pPr = _OxmlEl("w:pPr")
                    if style == "bullet" and bullet_numid:
                        numPr = _OxmlEl("w:numPr")
                        ilvl = _OxmlEl("w:ilvl")
                        ilvl.set(_qn("w:val"), "0")
                        numId_el = _OxmlEl("w:numId")
                        numId_el.set(_qn("w:val"), bullet_numid)
                        numPr.append(ilvl)
                        numPr.append(numId_el)
                        pPr.append(numPr)
                    new_p.append(pPr)

                    # Run
                    run_text = text
                    if style.startswith("number:"):
                        n = style.split(":", 1)[1]
                        run_text = f"{n}. {text}"

                    r = _OxmlEl("w:r")
                    if orig_rpr is not None:
                        r.append(copy.deepcopy(orig_rpr))
                    t = _OxmlEl("w:t")
                    t.text = run_text
                    t.set(_XML_SPACE, "preserve")
                    r.append(t)
                    new_p.append(r)

                    tc.insert(insert_idx + offset, new_p)

                replaced.add(matched)
                break  # one sentinel per tc

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _field_default(field: dict) -> str:
    """Return the configured default for a field as a string (empty string if absent)."""
    d = field.get("default")
    if d is None:
        return ""
    if isinstance(d, bool):
        return CHECKBOX_CHAR if d else UNCHECKED_CHAR
    return str(d)


def _count_visible_fields(section: dict) -> int:
    """Count configured fields that should be shown in a form section."""
    return sum(
        1 for field in section.get("fields", [])
        if field.get("placeholder") and not field.get("hidden")
    )


def _sanitize_non_negative_integer(value: Any) -> str:
    """Extract a non-negative integer from a price value, or return blank."""
    text = str(value or "").strip()
    if re.search(r"-\s*\d", text):
        return ""
    match = re.search(r"(?<![\d.])\d[\d,]*(?![\d.])", text)
    if not match:
        return ""
    digits = match.group(0).replace(",", "")
    return digits


def _ensure_chinese_language(language_key: str) -> None:
    """Keep 中文 selected whenever 繁體 or 簡體 is selected."""
    selected = st.session_state.get(language_key, [])
    if any(label in selected for label in ("繁體", "簡體", "简体")) and "中文" not in selected:
        st.session_state[language_key] = ["中文", *selected]


def _ensure_chinese_checkbox(
    chinese_key: str,
    traditional_key: str,
    simplified_key: str,
) -> None:
    """Keep the 中文 checkbox selected for traditional or simplified Chinese."""
    if st.session_state.get(traditional_key) or st.session_state.get(simplified_key):
        st.session_state[chinese_key] = True


def build_initial_form_values(extracted_flat: dict[str, str]) -> dict[str, str]:
    """Map extracted flat JSON keys → template placeholder names via field_config.
    Falls back to the field's ``default`` value when no AI-extracted value is available.
    """
    trim_size = extracted_flat.get("trim_size", "")
    raw_language = extracted_flat.get("language", "")
    lang_str = ", ".join(raw_language) if isinstance(
        raw_language, list) else str(raw_language)
    normalized_lang_str = lang_str.replace("简体", "簡體")

    # Parse both trim dimensions, then assign larger → 長 (#trim_width) and
    # smaller → 闊 (#trim_height), regardless of order in the source string.
    _raw0 = _parse_trim_dimension(trim_size, 0)
    _raw1 = _parse_trim_dimension(trim_size, 1)
    try:
        _d0, _d1 = float(_raw0 or 0), float(_raw1 or 0)
        _trim_長 = _mm_to_cm(str(max(_d0, _d1))) if max(_d0, _d1) else ""
        _trim_闊 = _mm_to_cm(str(min(_d0, _d1))) if min(_d0, _d1) else ""
    except (ValueError, TypeError):
        _trim_長 = _mm_to_cm(_raw0) if _raw0 else ""
        _trim_闊 = _mm_to_cm(_raw1) if _raw1 else ""

    values: dict[str, str] = {}
    for section in SECTIONS:
        for field in section["fields"]:
            ph = field["placeholder"]
            if not ph:
                continue
            src = field.get("source_key", "")
            if src == "#trim_width":
                values[ph] = _trim_長 or _field_default(field)
            elif src == "#trim_height":
                values[ph] = _trim_闊 or _field_default(field)
            elif src.startswith("#lang:"):
                keyword = src[len("#lang:"):]
                language_selected = keyword in normalized_lang_str
                if keyword == "中文":
                    language_selected = language_selected or any(
                        label in normalized_lang_str for label in ("繁體", "簡體")
                    )
                values[ph] = CHECKBOX_CHAR if language_selected else UNCHECKED_CHAR
            elif src and src in extracted_flat:
                raw = extracted_flat[src]
                if field.get("type") == "positive_integer":
                    raw = _sanitize_non_negative_integer(raw)
                elif ph == "publishDate":
                    raw = _format_date_ddmmyy(raw)
                elif ph in _CM_PLACEHOLDERS and raw:
                    raw = _mm_to_cm(raw)
                values[ph] = raw or _field_default(field)
            else:
                values[ph] = _field_default(field)
    return values


def _collect_current_values() -> dict[str, str]:
    """Read all form widget values from session state (read-only, no side effects)."""
    all_vals: dict[str, str] = {}
    for section in SECTIONS:
        for field in section["fields"]:
            ph = field["placeholder"]
            if not ph:
                continue
            key = f"form_{ph}"
            if field.get("hidden"):
                all_vals[ph] = _field_default(field)
                continue
            if field["type"] == "option":
                all_vals[ph] = field["label"]
                continue
            if field["type"] == "checkbox":
                val = st.session_state.get(key, False)
                all_vals[ph] = CHECKBOX_CHAR if val else UNCHECKED_CHAR
            elif field["type"] == "date":
                _dval = st.session_state.get(key)
                if hasattr(_dval, "strftime"):
                    all_vals[ph] = _dval.strftime("%d/%m/%y")
                else:
                    all_vals[ph] = _format_date_ddmmyy(
                        str(_dval)) if _dval else ""
            elif field["type"] == "positive_integer":
                all_vals[ph] = _sanitize_non_negative_integer(
                    st.session_state.get(key, ""))
            elif field["type"] == "html":
                # st_quill returns None while its iframe is still loading.
                # _quill_bak_<ph> stores the last confirmed non-None value so
                # the DOCX stays populated across reruns even when quill resets.
                val = st.session_state.get(key)
                if val is None:
                    val = st.session_state.get(
                        f"_quill_bak_{ph}",
                        st.session_state.get(
                            "form_values", {}).get(ph, "") or "",
                    )
                all_vals[ph] = val
            else:
                all_vals[ph] = st.session_state.get(key) or ""
    return all_vals


def _clear_form_widget_state() -> None:
    """Remove widget-key session entries so the form re-renders with fresh values."""
    option_keys = {
        f"form_option_{group_name}"
        for group_name in {
            field["option_group"]
            for section in SECTIONS
            for field in section["fields"]
            if field.get("type") == "option"
            and field.get("option_mode") == "single"
        }
    }
    for key in option_keys:
        st.session_state.pop(key, None)
    for section in SECTIONS:
        for field in section["fields"]:
            if not field["placeholder"]:
                continue
            key = f"form_{field['placeholder']}"
            if key in st.session_state:
                del st.session_state[key]
            if field.get("type") == "html":
                bak = f"_quill_bak_{field['placeholder']}"
                if bak in st.session_state:
                    del st.session_state[bak]
            if field.get("type") == "option":
                st.session_state.pop(key, None)
    for group_name in {
        field["option_group"]
        for section in SECTIONS
        for field in section["fields"]
        if field.get("type") == "option"
    }:
        st.session_state.pop(f"_option_values_{group_name}", None)


# ---------------------------------------------------------------------------
# Template helpers
# ---------------------------------------------------------------------------

def render_docx_template(template_path: str, context: dict[str, str]) -> bytes:
    tpl = DocxTemplate(template_path)
    tpl.render(context)
    buf = io.BytesIO()
    tpl.save(buf)
    buf.seek(0)
    return buf.read()


def _apply_option_strikethrough(
    doc_bytes: bytes,
    option_selections: dict[str, list[str]],
) -> bytes:
    """Double-strike unselected visible options in the v2 template."""
    if not option_selections:
        return doc_bytes

    doc = Document(io.BytesIO(doc_bytes))
    from docx.oxml import OxmlElement
    w_p = qn("w:p")
    w_r = qn("w:r")
    w_rpr = qn("w:rPr")
    w_t = qn("w:t")
    w_dstrike = qn("w:dstrike")
    option_groups: dict[str, list[dict]] = {}
    for section in SECTIONS:
        for field in section["fields"]:
            if field.get("type") == "option":
                option_groups.setdefault(
                    field["option_group"], []).append(field)
    for group_name, selected in option_selections.items():
        group_fields = option_groups.get(group_name)
        if not group_fields:
            continue
        labels = {field["label"]: field["option_value"]
                  for field in group_fields if field.get("option_value")}
        for para in doc.part.element.iter(w_p):
            para_text = "".join(node.text or "" for node in para.iter(w_t))
            if sum(label in para_text for label in labels) < 2:
                continue
            for run in para.iter(w_r):
                label = "".join(
                    node.text or "" for node in run.iter(w_t)).strip()
                if label in labels:
                    rpr = run.find(w_rpr)
                    if rpr is None:
                        rpr = OxmlElement("w:rPr")
                        run.insert(0, rpr)
                    existing = rpr.find(w_dstrike)
                    if labels[label] not in selected:
                        if existing is None:
                            rpr.append(OxmlElement("w:dstrike"))
                    elif existing is not None:
                        rpr.remove(existing)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


@st.cache_data(show_spinner=False)
def _render_docx_cached(template_path: str, context_json: str) -> bytes:
    """Generate DOCX bytes, cached by template path + JSON context."""
    return render_docx_template(template_path, json.loads(context_json))


# ---------------------------------------------------------------------------
# Page
# ---------------------------------------------------------------------------
def main() -> None:
    enforce_workspace_auth()

    st.title("📝 紅出版 新書表單產生器")

    with st.sidebar:
        st.header("⚙️ 設定")
        st.caption(f"已登入：{getattr(st.user, 'email', '未知帳號')}")
        if st.button("登出", key="sidebar_logout"):
            st.logout()
            st.stop()

    if "form_values" not in st.session_state:
        st.session_state.form_values = {}

    if not TEMPLATE_FILE.exists():
        st.error("專案根目錄找不到 template-v2.docx。")
        st.stop()

    try:
        prompt_text = read_prompt_text()
    except Exception as exc:
        st.error(str(exc))
        st.stop()

    vertex_settings = st.secrets.get("google_vertex", {})
    if not vertex_settings.get("project_id") or not vertex_settings.get("location"):
        st.error(
            "在 .streamlit/secrets.toml 的 [google_vertex] 中找不到 "
            "project_id 或 location。"
        )
        st.stop()
    vertex_client = None

    uploaded_file = st.file_uploader(
        "上傳新書Media Letter Word檔 (docx)，轉換成新書資料表格。", type=["docx"]
    )
    extract_clicked = st.button("提交", type="primary")

    if extract_clicked:
        if uploaded_file is None:
            st.warning("請先上傳 DOCX 檔案。")
            st.stop()

        with st.spinner("正在讀取上傳的 DOCX..."):
            input_bytes = uploaded_file.getvalue()
            document_text = extract_text_from_docx(input_bytes)

        if not document_text.strip():
            st.error("無法從上傳 DOCX 擷取可讀文字。")
            st.stop()

        composed_prompt = compose_vertex_prompt(prompt_text, document_text)

        start_time = time.monotonic()

        with st.spinner("正在透過AI擷取文件內容..."):
            try:
                if vertex_client is None:
                    vertex_client = get_google_vertex_client()
                raw_json, raw_response = call_google_vertex(
                    vertex_client, VERTEX_MODEL, composed_prompt
                )
                elapsed = time.monotonic() - start_time
                add_run(
                    file_name=uploaded_file.name,
                    model=VERTEX_MODEL,
                    raw_prompt=composed_prompt,
                    raw_response=raw_response,
                    elapsed_seconds=elapsed,
                    success=True,
                    user_email=getattr(st.user, "email", ""),
                    source_file=input_bytes,
                )
                normalized = normalize_extracted_payload(raw_json)
                extracted_flat = flatten_extracted_data(normalized)
                _clear_form_widget_state()
                st.session_state.form_values = build_initial_form_values(
                    extracted_flat)
                st.success("擷取完成，您可在下方檢查並編輯欄位。")
            except ValidationError as exc:
                elapsed = time.monotonic() - start_time
                add_run(
                    file_name=uploaded_file.name,
                    model=VERTEX_MODEL,
                    raw_prompt=composed_prompt,
                    raw_response="",
                    elapsed_seconds=elapsed,
                    success=False,
                    error_message=str(exc),
                    user_email=getattr(st.user, "email", ""),
                    source_file=input_bytes,
                )
                st.error("Google Vertex 回傳的資料格式不符合預期。")
                st.code(str(exc))
                st.stop()
            except Exception as exc:
                elapsed = time.monotonic() - start_time
                add_run(
                    file_name=getattr(uploaded_file, "name", "unknown"),
                    model=VERTEX_MODEL,
                    raw_prompt=composed_prompt,
                    raw_response="",
                    elapsed_seconds=elapsed,
                    success=False,
                    error_message=str(exc),
                    user_email=getattr(st.user, "email", ""),
                    source_file=input_bytes,
                )
                st.error(f"資料擷取失敗：{exc}")
                st.stop()

    if st.session_state.form_values:
        st.markdown("---")
        st.caption("AI資料擷取不一定完全正確，請務必檢查並補充必要欄位後再下載表格。")

        # ── Render all sections in order (html fields inline at correct position) ──
        _CHECKBOX_COLS = 4
        fv = st.session_state.form_values  # shorthand

        for section in SECTIONS:
            fields = section["fields"]
            layout = section.get("layout", "grid")

            if _count_visible_fields(section) > 1:
                st.subheader(section["title"])

            if layout == "options":
                option_fields = [f for f in fields if f["type"] == "option"]
                option_labels = [f["label"] for f in option_fields]
                option_values = {f["label"]: f["option_value"]
                                 for f in option_fields}
                option_key = f"form_option_{section['option_group']}"
                initial = [
                    field["label"] for field in option_fields
                    if fv.get(field["placeholder"]) == CHECKBOX_CHAR
                ]
                if section.get("option_mode") == "multi":
                    selected_labels = st.multiselect(
                        section["title"], options=option_labels,
                        default=initial, key=option_key)
                else:
                    selected_label = st.selectbox(
                        section["title"], options=[""] + option_labels,
                        index=(option_labels.index(
                            initial[0]) + 1) if initial else 0,
                        key=option_key)
                    selected_labels = [
                        selected_label] if selected_label else []
                st.session_state[f"_option_values_{section['option_group']}"] = [
                    option_values[label] for label in selected_labels
                ]
                for field in fields:
                    if field["type"] == "text":
                        ph = field["placeholder"]
                        st.text_input(field["label"], value=fv.get(
                            ph, ""), key=f"form_{ph}")
            elif layout == "grid":
                # Grid layout: row_group fields share a row; ungrouped fields pair 2-per-row
                others = [f for f in fields if f["type"]
                          not in ("checkbox", "option") and not f.get("hidden")]
                checkboxes = [f for f in fields if f["type"]
                              == "checkbox" and not f.get("hidden")]
                embedded_options = [f for f in fields if f["type"] == "option"]

                # Build row batches
                _row_batches: list[list[dict]] = []
                _ri = 0
                while _ri < len(others):
                    _rg = others[_ri].get("row_group")
                    if _rg:
                        _grp = [others[_ri]]
                        _rj = _ri + 1
                        while _rj < len(others) and others[_rj].get("row_group") == _rg:
                            _grp.append(others[_rj])
                            _rj += 1
                        _row_batches.append(_grp)
                        _ri = _rj
                    else:
                        if _ri + 1 < len(others) and not others[_ri + 1].get("row_group"):
                            _row_batches.append([others[_ri], others[_ri + 1]])
                            _ri += 2
                        else:
                            _row_batches.append([others[_ri]])
                            _ri += 1

                for _row_fields in _row_batches:
                    cols = st.columns(len(_row_fields))
                    for col, field in zip(cols, _row_fields):
                        ph = field["placeholder"]
                        key = f"form_{ph}"
                        hint = field.get("hint") or None
                        opts = field.get("options")
                        with col:
                            if field["type"] == "positive_integer":
                                if key not in st.session_state:
                                    initial = _sanitize_non_negative_integer(
                                        fv.get(ph, ""))
                                    st.number_input(
                                        field["label"], min_value=0, step=1,
                                        value=int(
                                            initial) if initial else None,
                                        format="%d", key=key,
                                    )
                                else:
                                    st.number_input(
                                        field["label"], min_value=0, step=1,
                                        format="%d", key=key,
                                    )
                            elif field["type"] == "html":
                                st.markdown(f"**{field['label']}**")
                                _qval = st.session_state.get(key)
                                if _qval is not None:
                                    st.session_state[f"_quill_bak_{ph}"] = _qval
                                _qinit = _qval if _qval is not None else fv.get(
                                    ph, "")
                                st_quill(value=_qinit, html=True, key=key)
                            elif field["type"] == "date":
                                if key not in st.session_state:
                                    st.date_input(field["label"], value=_parse_to_date(
                                        fv.get(ph, "")), key=key, format="DD/MM/YYYY")
                                else:
                                    st.date_input(
                                        field["label"], key=key, format="DD/MM/YYYY")
                            elif opts:
                                current = fv.get(ph, "")
                                idx = opts.index(
                                    current) if current in opts else 0
                                st.selectbox(
                                    field["label"], options=opts, index=idx, key=key)
                            elif field["type"] == "textarea":
                                if key not in st.session_state:
                                    st.text_area(field["label"], value=fv.get(
                                        ph, ""), key=key, placeholder=hint)
                                else:
                                    st.text_area(
                                        field["label"], key=key, placeholder=hint)
                            else:
                                if key not in st.session_state:
                                    st.text_input(field["label"], value=fv.get(
                                        ph, ""), key=key, placeholder=hint)
                                else:
                                    st.text_input(
                                        field["label"], key=key, placeholder=hint)
                for i in range(0, len(checkboxes), _CHECKBOX_COLS):
                    group = checkboxes[i:i + _CHECKBOX_COLS]
                    cols = st.columns(_CHECKBOX_COLS)
                    for j, field in enumerate(group):
                        ph = field["placeholder"]
                        key = f"form_{ph}"
                        with cols[j]:
                            if key not in st.session_state:
                                default_val = fv.get(
                                    ph, _field_default(field)) == CHECKBOX_CHAR
                                st.checkbox(field["label"],
                                            value=default_val, key=key)
                            else:
                                st.checkbox(field["label"], key=key)

                option_groups: dict[str, list[dict]] = {}
                for field in embedded_options:
                    option_groups.setdefault(
                        field["option_group"], []).append(field)
                for group_name, group_fields in option_groups.items():
                    group_title = {
                        "edition": "版次",
                        "sales_level": "銷售級別",
                        "language": "語種",
                    }.get(group_name, group_name)
                    option_values = {
                        field["label"]: field["option_value"] for field in group_fields
                    }
                    if group_fields[0].get("option_mode") == "multi":
                        option_labels = [field["label"]
                                         for field in group_fields]
                        initial = [
                            field["label"] for field in group_fields
                            if fv.get(field["placeholder"]) == CHECKBOX_CHAR
                        ]
                        selected_labels = st.multiselect(
                            group_title,
                            options=option_labels,
                            default=initial,
                            key=f"form_option_{group_name}",
                            on_change=_ensure_chinese_language,
                            args=(f"form_option_{group_name}",),
                        )
                    else:
                        option_key = f"form_option_{group_name}"
                        initial = next(
                            (field["label"] for field in group_fields
                             if fv.get(field["placeholder"]) == CHECKBOX_CHAR),
                            group_fields[0]["label"],
                        )
                        selected_label = st.selectbox(
                            group_title,
                            options=[field["label"] for field in group_fields],
                            index=[field["label"]
                                   for field in group_fields].index(initial),
                            key=option_key,
                        )
                        selected_labels = [selected_label]
                    st.session_state[f"_option_values_{group_name}"] = [
                        option_values[label] for label in selected_labels
                    ]
            elif layout == "row3":
                # 3 fields per row: publisher | title | price (competitor rows)
                visible = [f for f in fields if f.get(
                    "placeholder") and not f.get("hidden")]
                for i in range(0, len(visible), 3):
                    group = visible[i:i + 3]
                    cols = st.columns(len(group))
                    for col, field in zip(cols, group):
                        ph = field["placeholder"]
                        key = f"form_{ph}"
                        hint = field.get("hint") or None
                        with col:
                            if key not in st.session_state:
                                st.text_input(field["label"], value=fv.get(ph, ""),
                                              key=key, placeholder=hint)
                            else:
                                st.text_input(
                                    field["label"], key=key, placeholder=hint)
            else:
                # Render in config order: group consecutive checkboxes 4/row,
                # non-checkbox fields rendered individually at their config position
                idx = 0
                while idx < len(fields):
                    field = fields[idx]
                    if field["type"] == "checkbox":
                        # Collect run of consecutive checkboxes
                        run: list[dict] = []
                        while idx < len(fields) and fields[idx]["type"] == "checkbox":
                            run.append(fields[idx])
                            idx += 1
                        for j in range(0, len(run), _CHECKBOX_COLS):
                            group = run[j:j + _CHECKBOX_COLS]
                            cols = st.columns(_CHECKBOX_COLS)
                            for k, cb_field in enumerate(group):
                                ph = cb_field["placeholder"]
                                key = f"form_{ph}"
                                with cols[k]:
                                    if key not in st.session_state:
                                        default_val = fv.get(
                                            ph, _field_default(cb_field)) == CHECKBOX_CHAR
                                        st.checkbox(
                                            cb_field["label"], value=default_val, key=key)
                                    else:
                                        st.checkbox(cb_field["label"], key=key)
                    else:
                        if field["type"] == "label":
                            st.markdown(f"**{field['label']}**")
                            idx += 1
                        elif field.get("row_group"):
                            _rg = field["row_group"]
                            _rg_fields = [field]
                            _look = idx + 1
                            while _look < len(fields) and fields[_look].get("row_group") == _rg:
                                _rg_fields.append(fields[_look])
                                _look += 1
                            _rg_cols = st.columns(len(_rg_fields))
                            for _rg_col, _rg_field in zip(_rg_cols, _rg_fields):
                                _ph = _rg_field["placeholder"]
                                _key = f"form_{_ph}"
                                _hint = _rg_field.get("hint") or None
                                _opts = _rg_field.get("options")
                                with _rg_col:
                                    if _rg_field["type"] == "positive_integer":
                                        if _key not in st.session_state:
                                            _initial = _sanitize_non_negative_integer(
                                                fv.get(_ph, ""))
                                            st.number_input(
                                                _rg_field["label"], min_value=0, step=1,
                                                value=int(
                                                    _initial) if _initial else None,
                                                format="%d", key=_key,
                                            )
                                        else:
                                            st.number_input(
                                                _rg_field["label"], min_value=0, step=1,
                                                format="%d", key=_key,
                                            )
                                    elif _rg_field["type"] == "date":
                                        if _key not in st.session_state:
                                            st.date_input(_rg_field["label"], value=_parse_to_date(
                                                fv.get(_ph, "")), key=_key, format="DD/MM/YYYY")
                                        else:
                                            st.date_input(
                                                _rg_field["label"], key=_key, format="DD/MM/YYYY")
                                    elif _opts:
                                        _cur = fv.get(_ph, "")
                                        _sidx = _opts.index(
                                            _cur) if _cur in _opts else 0
                                        st.selectbox(
                                            _rg_field["label"], options=_opts, index=_sidx, key=_key)
                                    elif _rg_field["type"] == "textarea":
                                        if _key not in st.session_state:
                                            st.text_area(_rg_field["label"], value=fv.get(
                                                _ph, ""), key=_key, placeholder=_hint)
                                        else:
                                            st.text_area(
                                                _rg_field["label"], key=_key, placeholder=_hint)
                                    else:
                                        if _key not in st.session_state:
                                            st.text_input(_rg_field["label"], value=fv.get(
                                                _ph, ""), key=_key, placeholder=_hint)
                                        else:
                                            st.text_input(
                                                _rg_field["label"], key=_key, placeholder=_hint)
                            idx = _look
                        else:
                            ph = field["placeholder"]
                            key = f"form_{ph}"
                            hint = field.get("hint") or None
                            opts = field.get("options")
                            if field["type"] == "positive_integer":
                                if key not in st.session_state:
                                    initial = _sanitize_non_negative_integer(
                                        fv.get(ph, ""))
                                    st.number_input(
                                        field["label"], min_value=0, step=1,
                                        value=int(
                                            initial) if initial else None,
                                        format="%d", key=key,
                                    )
                                else:
                                    st.number_input(
                                        field["label"], min_value=0, step=1,
                                        format="%d", key=key,
                                    )
                            elif field["type"] == "html":
                                st.markdown(f"**{field['label']}**")
                                _qval = st.session_state.get(key)
                                if _qval is not None:
                                    st.session_state[f"_quill_bak_{ph}"] = _qval
                                _qinit = _qval if _qval is not None else fv.get(
                                    ph, "")
                                st_quill(value=_qinit, html=True, key=key)
                            elif field["type"] == "date":
                                if key not in st.session_state:
                                    st.date_input(field["label"], value=_parse_to_date(
                                        fv.get(ph, "")), key=key, format="DD/MM/YYYY")
                                else:
                                    st.date_input(
                                        field["label"], key=key, format="DD/MM/YYYY")
                            elif opts:
                                current = fv.get(ph, "")
                                sidx = opts.index(
                                    current) if current in opts else 0
                                st.selectbox(
                                    field["label"], options=opts, index=sidx, key=key)
                            elif field["type"] == "textarea":
                                if key not in st.session_state:
                                    st.text_area(field["label"], value=fv.get(
                                        ph, ""), key=key, placeholder=hint)
                                else:
                                    st.text_area(
                                        field["label"], key=key, placeholder=hint)
                            else:
                                if key not in st.session_state:
                                    st.text_input(field["label"], value=fv.get(
                                        ph, ""), key=key, placeholder=hint)
                                else:
                                    st.text_input(
                                        field["label"], key=key, placeholder=hint)
                            idx += 1

            st.markdown("---")

        _current_vals = _collect_current_values()
        _docx_ctx = dict(_current_vals)

        # Save original HTML for bullet fields before any conversion
        _orig_html_for_bullets = {
            _ph: _docx_ctx.get(_ph, "")
            for _ph in _BULLET_SENTINEL
        }

        # Convert HTML (from st_quill) to plain text for all html-type fields
        for _sec in SECTIONS:
            for _f in _sec["fields"]:
                if _f.get("type") == "html" and _f.get("placeholder"):
                    _ph = _f["placeholder"]
                    if _docx_ctx.get(_ph):
                        _docx_ctx[_ph] = _html_to_plain_text(_docx_ctx[_ph])

        # For bullet fields: parse HTML into items and insert sentinels so that
        # the post-processor can replace them with native Word bullet paragraphs.
        _bullet_sentinel_items: dict[str, list[tuple[str, str]]] = {}
        for _ph, _sentinel in _BULLET_SENTINEL.items():
            _orig = _orig_html_for_bullets.get(_ph, "")
            if _orig:
                _items = _parse_html_to_items(_orig)
                if _items:
                    _docx_ctx[_ph] = _sentinel
                    _bullet_sentinel_items[_sentinel] = _items

        _docx_bytes = _render_docx_cached(
            str(TEMPLATE_FILE), json.dumps(
                _docx_ctx, sort_keys=True, ensure_ascii=False)
        )

        _docx_bytes = _apply_option_strikethrough(
            _docx_bytes,
            {
                group_name: st.session_state.get(
                    f"_option_values_{group_name}", [])
                for group_name in {
                    field["option_group"]
                    for section in SECTIONS
                    for field in section["fields"]
                    if field.get("type") == "option"
                }
            },
        )

        # Post-process: replace sentinels with native Word bullet paragraphs
        if _bullet_sentinel_items:
            _docx_bytes = _apply_bullet_formatting(
                _docx_bytes, _bullet_sentinel_items)
        st.download_button(
            label="⬇️ 下載表格",
            data=_docx_bytes,
            file_name="新書資料表格.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            use_container_width=True,
        )


if __name__ == "__main__":
    main()
